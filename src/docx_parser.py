"""
DOCX Parser for Microsoft Word Documents
Week 3 Day 17 - Extract text and metadata from .docx files
"""

from typing import Dict, Optional, List
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. Run: pip install python-docx")


class DOCXParser:
    """Extract text and metadata from Word documents"""

    def __init__(self):
        self.stats = {
            "files_processed": 0,
            "paragraphs_extracted": 0,
            "tables_extracted": 0,
        }

        print(f"📝 DOCX Parser initialized")
        print(f"   Available: {DOCX_AVAILABLE}")

    def parse_file(self, filepath: str) -> Optional[Dict]:
        """Parse a DOCX file and extract content + metadata"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx not available")
            return None

        filepath = Path(filepath)

        if not filepath.exists():
            print(f"❌ File not found: {filepath}")
            return None

        if filepath.suffix.lower() != ".docx":
            print(f"❌ Not a DOCX file: {filepath}")
            return None

        print(f"\n📝 Parsing DOCX: {filepath.name}")

        try:
            # Load document
            doc = Document(filepath)

            # Extract content with structure
            content = self._extract_content_with_structure(doc)

            # Extract metadata
            metadata = self._extract_metadata(doc, filepath)

            # Statistics
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            word_count = len(content.split())

            result = {
                "filename": filepath.name,
                "filepath": str(filepath),
                "content": content,
                "metadata": metadata,
                "paragraph_count": para_count,
                "table_count": table_count,
                "word_count": word_count,
                "file_type": "docx",
            }

            print(
                f"   ✓ Extracted: {word_count} words, {para_count} paragraphs, {table_count} tables"
            )

            self.stats["files_processed"] += 1
            self.stats["paragraphs_extracted"] += para_count
            self.stats["tables_extracted"] += table_count

            return result

        except Exception as e:
            print(f"❌ Failed to parse {filepath.name}: {e}")
            return None

    def _extract_content_with_structure(self, doc) -> str:
        """Extract text preserving document structure (headings, lists, tables)"""
        content_parts = []

        # Iterate through document elements in order
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # It's a paragraph
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()

                if not text:
                    continue

                # Check if it's a heading
                if paragraph.style.name.startswith("Heading"):
                    level = self._get_heading_level(paragraph.style.name)
                    # Add markdown-style heading
                    content_parts.append(f"\n{'#' * level} {text}\n")
                else:
                    content_parts.append(text)

            elif isinstance(element, CT_Tbl):
                # It's a table
                table = Table(element, doc)
                table_text = self._extract_table_text(table)
                content_parts.append(f"\n[TABLE]\n{table_text}\n")

        return "\n\n".join(content_parts)

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name (e.g., 'Heading 1' -> 1)"""
        try:
            # Extract number from style name
            import re

            match = re.search(r"Heading (\d+)", style_name)
            if match:
                return int(match.group(1))
            return 1
        except:
            return 1

    def _extract_table_text(self, table) -> str:
        """Convert table to readable text format"""
        rows = []

        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)

            # Join cells with pipes (markdown table style)
            rows.append(" | ".join(cells))

        return "\n".join(rows)

    def _extract_metadata(self, doc, filepath: Path) -> Dict:
        """Extract document metadata"""
        metadata = {
            "filename": filepath.name,
            "source": "docx",
            "title": filepath.stem.replace("-", " ").replace("_", " ").title(),
        }

        try:
            # Access core properties
            core_props = doc.core_properties

            # Title
            if core_props.title:
                metadata["title"] = core_props.title

            # Author
            if core_props.author:
                metadata["author"] = core_props.author

            # Subject
            if core_props.subject:
                metadata["subject"] = core_props.subject

            # Keywords/Tags
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords

            # Dates
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()

            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()

            # Last modified by
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by

            # Revision
            if core_props.revision:
                metadata["revision"] = str(core_props.revision)

        except Exception as e:
            print(f"   ℹ️  Could not extract all metadata: {e}")

        return metadata

    def parse_folder(self, folder_path: str) -> List[Dict]:
        """Parse all DOCX files in a folder"""
        folder = Path(folder_path)

        if not folder.exists():
            print(f"❌ Folder not found: {folder_path}")
            return []

        docx_files = list(folder.glob("**/*.docx"))

        # Filter out temporary files (starting with ~$)
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]

        if not docx_files:
            print(f"⚠️  No DOCX files found in {folder_path}")
            return []

        print(f"\n📂 Found {len(docx_files)} DOCX file(s)")

        results = []
        for docx_file in docx_files:
            result = self.parse_file(docx_file)
            if result:
                results.append(result)

        return results

    def extract_headings_only(self, filepath: str) -> List[str]:
        """Extract only headings (useful for document outline)"""
        if not DOCX_AVAILABLE:
            return []

        try:
            doc = Document(filepath)
            headings = []

            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading"):
                    level = self._get_heading_level(paragraph.style.name)
                    indent = "  " * (level - 1)
                    headings.append(f"{indent}• {paragraph.text}")

            return headings
        except:
            return []

    def get_stats(self) -> Dict:
        """Get parsing statistics"""
        return {
            **self.stats,
            "avg_paragraphs_per_file": (
                self.stats["paragraphs_extracted"] / self.stats["files_processed"]
                if self.stats["files_processed"] > 0
                else 0
            ),
        }

    def print_stats(self):
        """Print parsing statistics"""
        stats = self.get_stats()

        print(f"\n📊 DOCX Parser Statistics:")
        print(f"   Files processed: {stats['files_processed']}")
        print(f"   Paragraphs extracted: {stats['paragraphs_extracted']}")
        print(f"   Tables extracted: {stats['tables_extracted']}")
        print(f"   Avg paragraphs/file: {stats['avg_paragraphs_per_file']:.1f}")


# Quick test
if __name__ == "__main__":
    print("=" * 70)
    print("📝 DOCX PARSER TEST")
    print("=" * 70)

    parser = DOCXParser()

    if not DOCX_AVAILABLE:
        print("\n❌ Install python-docx first:")
        print("   pip install python-docx")
    else:
        # Test with sample DOCX if available
        import os

        test_folder = "data/docx"

        if not os.path.exists(test_folder):
            os.makedirs(test_folder, exist_ok=True)
            print(f"\n📁 Created folder: {test_folder}")
            print(f"   Add some .docx files here to test!")
            print(f"\n💡 You can add:")
            print(f"   • Your resume/CV")
            print(f"   • Project notes")
            print(f"   • Meeting notes")
            print(f"   • Learning journals")
            print(f"   • Any Word document")
        else:
            # Parse all DOCX files in folder
            results = parser.parse_folder(test_folder)

            if results:
                print(f"\n✅ Parsed {len(results)} DOCX file(s)")

                # Show first result preview
                sample = results[0]
                print(f"\n--- Sample: {sample['filename']} ---")
                print(f"Title: {sample['metadata'].get('title', 'N/A')}")
                print(f"Author: {sample['metadata'].get('author', 'N/A')}")
                print(f"Paragraphs: {sample['paragraph_count']}")
                print(f"Tables: {sample['table_count']}")
                print(f"Words: {sample['word_count']}")
                print(f"\nContent preview:")
                print(sample["content"][:400] + "...")

                # Show document outline
                print(f"\n--- Document Outline ---")
                headings = parser.extract_headings_only(sample["filepath"])
                if headings:
                    for heading in headings:
                        print(heading)
                else:
                    print("   (No headings found)")

            parser.print_stats()

    print("\n" + "=" * 70)
    print("✅ DOCX PARSER READY!")
    print("=" * 70)
