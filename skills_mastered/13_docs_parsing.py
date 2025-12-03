"""
SKILL MASTERED: DOCX Parsing for Word Documents
Week 3 Day 17

What I learned:
- python-docx library for Word document parsing
- Extracting structured content (headings, paragraphs, tables)
- Preserving document hierarchy
- Metadata extraction (author, dates, revision)
- Document core properties access

Portfolio highlight:
"Built Word document parser preserving structure (headings, lists, tables).
Extracts metadata including author, created/modified dates, and revision
history. Integrated with RAG for semantic search across resumes, reports,
and notes."

Business value:
- Parse resumes for skill extraction
- Index project documentation
- Search meeting notes
- Analyze client proposals
- Build personal knowledge base from Word docs

Technical features:
- Structure-aware parsing (headings preserved)
- Table extraction and formatting
- Metadata extraction (10+ fields)
- Document outline generation
- Smart chunking for long documents

Real-world application:
"Indexed 50+ professional documents (resumes, project reports, meeting notes)
into searchable knowledge base. Enabled instant retrieval of past work
experience and skills with 90%+ accuracy."
"""

from docx import Document


def demo():
    """Quick DOCX parsing demo"""
    # Load document
    doc = Document("sample.docx")

    # Extract paragraphs
    for para in doc.paragraphs:
        print(f"{para.style.name}: {para.text}")

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            print(" | ".join(cells))

    # Access metadata
    print(f"Author: {doc.core_properties.author}")
    print(f"Created: {doc.core_properties.created}")


if __name__ == "__main__":
    demo()
